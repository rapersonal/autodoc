from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
import oracledb


# Init  database connection
oracledb.init_oracle_client(lib_dir=r"C:\oracle\instantclient_21_13")
# set the database values
db_user = ''
db_password = ''
db_encrypted_password = b''
db_host = ''
db_port = ''
db_service_name = ''

# Connect to the  database
connection = oracledb.connect(f'{db_user}/{db_password}@{db_host}:{db_port}/{db_service_name}')
cursor = connection.cursor()

# Execute a query to retrieve data from the database
query = """
SELECT 
    NAME,
    DESCRIPTION,
    HEADER,
    FOOTER,
    PARAGRAPH
FROM DOC_DATA A
"""
cursor.execute(query, [file_name])
results = cursor.fetchall()
# Load the Word document template
doc = Document('template.docx')
doc2 = Document('template6.docx')
# Clear all content in the document

copying = False

heading = doc2.add_heading(level=1)
run = heading.add_run('1.0 Introduction')
run.font.name = 'Roboto'
run.font.size = Pt(18)
run.bold = True
# Iterate over the paragraphs in the other document
for paragraph in doc.paragraphs:
    # If this is the start of the section to copy
    if '1.0 Introduction' in paragraph.text:
        copying = True
        skip_paragraph = True  # Skip the current paragraph
        continue  # Skip to the next iteration

    # If this is the end of the section to copy
    elif '2.0 Source Tables' in paragraph.text:
        copying = False
        break  # Stop iterating over the paragraphs

    # If we're in the section to copy
    if copying:
        # Add the paragraph to the new document and copy the paragraph style
            # Add the paragraph to the new document and copy the paragraph style
        new_paragraph = doc2.add_paragraph(style=paragraph.style.name)
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
           # Copy the run's formatting
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.color.rgb = run.font.color.rgb
            new_run.font.size = run.font.size
            new_run.style.name = run.style.name
            new_run.font.name = 'Roboto'
    # Copy the paragraph's spacing settings
        new_paragraph.paragraph_format.space_before = paragraph.paragraph_format.space_before
        new_paragraph.paragraph_format.space_after = paragraph.paragraph_format.space_after
        new_paragraph.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing


# Section 2 copy
heading = doc2.add_heading(level=1)
run = heading.add_run('2.0 Source Tables & Relations')
run.font.name = 'Roboto'
run.font.size = Pt(18)
run.bold = True
copying = False
# Iterate over the paragraphs in the other document
for paragraph in doc.paragraphs:
    # If this is the start of the section to copy
    if '2.0 Source Tables & Relations' in paragraph.text:
        copying = True
        skip_paragraph = True  # Skip the current paragraph
        continue  # Skip to the next iteration
    # If this is the end of the section to copy
    elif '3.0 Tables in replication layer' in paragraph.text:
        copying = False
        break  # Stop iterating over the paragraphs

    # If we're in the section to copy
    if copying:
        # Add the paragraph to the new document and copy the paragraph style
            # Add the paragraph to the new document and copy the paragraph style
        new_paragraph = doc2.add_paragraph(style=paragraph.style.name)
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
           # Copy the run's formatting
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.color.rgb = run.font.color.rgb
            new_run.font.size = run.font.size
            new_run.style.name = run.style.name
            new_run.font.name = 'Roboto'
    # Copy the paragraph's spacing settings
        new_paragraph.paragraph_format.space_before = paragraph.paragraph_format.space_before
        new_paragraph.paragraph_format.space_after = paragraph.paragraph_format.space_after
        new_paragraph.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing

# Section 3 copy
heading = doc2.add_heading(level=1)
run = heading.add_run('3.0 Tables in replication layer')
run.font.name = 'Roboto'
run.font.size = Pt(18)
run.bold = True
copying = False
copying = False
# Iterate over the paragraphs in the other document
for paragraph in doc.paragraphs:
    # If this is the start of the section to copy
    if '3.0 Tables in replication layer' in paragraph.text:
        copying = True
        skip_paragraph = True  # Skip the current paragraph
        continue  # Skip to the next iteration
    # If this is the end of the section to copy
    elif '4.0 Data Read and Load Process' in paragraph.text:
        copying = False
        break  # Stop iterating over the paragraphs

    # If we're in the section to copy
    if copying:
        # Add the paragraph to the new document and copy the paragraph style
            # Add the paragraph to the new document and copy the paragraph style
        new_paragraph = doc2.add_paragraph(style=paragraph.style.name)
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
           # Copy the run's formatting
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.color.rgb = run.font.color.rgb
            new_run.font.size = run.font.size
            new_run.style.name = run.style.name
            new_run.font.name = 'Roboto'
    # Copy the paragraph's spacing settings
        new_paragraph.paragraph_format.space_before = paragraph.paragraph_format.space_before
        new_paragraph.paragraph_format.space_after = paragraph.paragraph_format.space_after
        new_paragraph.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
# section 3 add table
# Iterate over the results and populate the document
# Find the "SCM Table Analysis" heading and replace the table below it
found_heading = False
for paragraph in doc.paragraphs:
    if 'SCM Table Analysis' in paragraph.text:
        doc2.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        print("Found the 'SCM Table Analysis' heading")  # Print a message in the console
        found_heading = True

    # If the heading has been found, delete the first following table
    if found_heading:
        for table in doc.tables:
            if table._element.getparent() is paragraph._element.getparent():
                # Delete the table
                table._element.getparent().remove(table._element)
                break  # Stop after the first table
        break  # Stop after the first occurrence

# Add a new table
table = doc2.add_table(rows=len(results) + 1, cols=6)  # +1 for the header row
# Set table style
table.style = 'Table Grid'

# Set cell margins and vertical alignment
for row in table.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.margin_top = Pt(5)
        cell.margin_bottom = Pt(5)
        
# Set header row background color, text color, and font size
for cell in table.rows[0].cells:
    cell.text = cell.text.upper()  # Convert text to uppercase
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # White text
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), '0F9ED5'))  # Blue background
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
# Define your color
header_font_color = RGBColor(0xFF, 0xFF, 0xFF)  # Convert FFFFFF from hex to RGB

# Add the header row
for j, header in enumerate(['EBS Table/ View Name', 'Schema Name', 'Access in ODS', 'Available in ODS', 'Rows (Approx)', 'Replication needed']):
    cell = table.cell(0, j)
    cell.text = header
    # Set font color
    run = cell.paragraphs[0].runs[0]
    run.font.color.rgb = header_font_color
# Populate the table with data
for i, row_data in enumerate(results, start=1):  # start=1 to skip the header row
    for j, cell_data in enumerate(row_data):
        table.cell(i, j).text = str(cell_data)

# Close the database connection
cursor.close()
connection.close()

# Section 4 copy
heading = doc2.add_heading(level=1)
run = heading.add_run('4.0 Data Read and Load Process')
run.font.name = 'Roboto'
run.font.size = Pt(18)
run.bold = True
copying = False
copying = False

copying = False
# Iterate over the paragraphs in the other document
for paragraph in doc.paragraphs:
    # If this is the start of the section to copy
    if '4.0 Data Read and Load Process' in paragraph.text:
        copying = True
        skip_paragraph = True  # Skip the current paragraph
        continue  # Skip to the next iteration
    # If this is the end of the section to copy
    elif '5.0 Appendices' in paragraph.text:
        copying = False
        break  # Stop iterating over the paragraphs

    # If we're in the section to copy
    if copying:
        # Add the paragraph to the new document and copy the paragraph style
            # Add the paragraph to the new document and copy the paragraph style
        new_paragraph = doc2.add_paragraph(style=paragraph.style.name)
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
           # Copy the run's formatting
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.color.rgb = run.font.color.rgb
            new_run.font.size = run.font.size
            new_run.style.name = run.style.name
            new_run.font.name = 'Roboto'
    # Copy the paragraph's spacing settings
        new_paragraph.paragraph_format.space_before = paragraph.paragraph_format.space_before
        new_paragraph.paragraph_format.space_after = paragraph.paragraph_format.space_after
        new_paragraph.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing


for paragraph in doc2.paragraphs:
    for run in paragraph.runs:
        if '<<FILE_NAME_TEXT>>' in run.text:
            run.text = run.text.replace('<<FILE_NAME_TEXT>>', file_name_text)

heading = doc2.add_heading(level=1)
run = heading.add_run('5.0 Appendices')
run.font.name = 'Roboto'
run.font.size = Pt(18)
run.bold = True


# Add a table of contents
doc2.add_page_break()
doc2.add_toc()

# Save the populated document
doc2.save(doc_name)

